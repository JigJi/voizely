"""Export MoM to DOCX with Appworks template formatting."""
import json
import os
import tempfile

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

LOGO_PATH = os.path.join(os.path.dirname(__file__), "..", "static", "img", "logo_1.png")
HEADER_COLOR = "1F4E79"  # Dark blue
LABEL_BG = "D9D9D9"  # Light gray


def _set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def _add_styled_para(doc, text, bold=False, size=18, align=None, color=None, space_after=4):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = "TH SarabunPSK"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def _style_table_cell(cell, text, bold=False, size=14, bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    run.font.name = "TH SarabunPSK"
    run.font.size = Pt(size)
    run.bold = bold
    if bg:
        _set_cell_shading(cell, bg)


def export_mom_docx(transcription, segments, db):
    """Generate DOCX file and return temp file path."""
    title = transcription.auto_title or transcription.audio_file.original_filename
    duration = segments[-1].end_time if segments else 0
    speakers = list(dict.fromkeys(s.speaker for s in segments))
    actions = json.loads(transcription.action_items) if transcription.action_items else []
    decisions = json.loads(transcription.key_decisions) if transcription.key_decisions else []

    # Parse MoM sections
    mom = transcription.mom_full or transcription.summary or ""
    sections = {}
    current = None
    for line in mom.split("\n"):
        if line.startswith("### ") and "ข้อมูลการประชุม" not in line:
            current = line.replace("### ", "").strip()
            sections[current] = []
        elif current and line.strip() and line.strip() != "---":
            sections.setdefault(current, []).append(line)

    doc = Document()

    # Page margins (match template)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(1.9)

    # Header with logo
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(LOGO_PATH):
        run = hp.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.5))

    # Style defaults
    style = doc.styles["Normal"]
    style.font.name = "TH SarabunPSK"
    style.font.size = Pt(14)

    # Title
    _add_styled_para(doc, "รายงานสรุปการประชุม", bold=True, size=20,
                     align=WD_ALIGN_PARAGRAPH.CENTER, color=HEADER_COLOR, space_after=12)

    # Info table
    info_data = [
        ("หัวข้อ", title),
        ("วันที่และเวลา", transcription.created_at.strftime("%d/%m/%Y %H:%M")),
        ("ความยาว", f"{int(duration//60)} นาที {int(duration%60)} วินาที"),
    ]
    table = doc.add_table(rows=len(info_data), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(info_data):
        _style_table_cell(table.rows[i].cells[0], label, bold=True, bg=LABEL_BG)
        _style_table_cell(table.rows[i].cells[1], value)
    # Set first column width
    for row in table.rows:
        row.cells[0].width = Cm(4)

    doc.add_paragraph()

    # ผู้เข้าร่วมประชุม
    _add_styled_para(doc, "ผู้เข้าร่วมประชุม", bold=True, size=16, color=HEADER_COLOR, space_after=4)
    for spk in speakers:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(spk)
        run.font.name = "TH SarabunPSK"
        run.font.size = Pt(14)

    doc.add_paragraph()

    # สรุปภาพรวม
    overview = sections.get("สรุปภาพรวม", [])
    if overview:
        _add_styled_para(doc, "สรุปภาพรวม", bold=True, size=16, color=HEADER_COLOR, space_after=4)
        _add_styled_para(doc, " ".join(overview), size=14, space_after=8)

    # ประเด็นที่พูดคุย (หัวข้อ + bullet เหมือน MoM web)
    topics_lines = sections.get("ประเด็นที่พูดคุย", [])
    if topics_lines:
        _add_styled_para(doc, "ข้อสรุปที่ประชุม", bold=True, size=16, color=HEADER_COLOR, space_after=4)

        for line in topics_lines:
            stripped = line.strip()
            if stripped.startswith("- **"):
                topic_name = stripped.replace("- **", "").replace("**", "").strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(topic_name)
                run.font.name = "TH SarabunPSK"
                run.font.size = Pt(14)
                run.bold = True
            elif stripped.startswith("- ") or stripped.startswith("  -"):
                bullet_text = stripped.lstrip("- ").strip()
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(bullet_text)
                run.font.name = "TH SarabunPSK"
                run.font.size = Pt(14)

        doc.add_paragraph()

    # มติที่ประชุม
    decisions_lines = sections.get("มติที่ประชุม", [])
    if decisions_lines:
        _add_styled_para(doc, "มติที่ประชุม", bold=True, size=16, color=HEADER_COLOR, space_after=4)
        for line in decisions_lines:
            text = line.strip().lstrip("- ").strip()
            if text and text != "ไม่มี":
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(text)
                run.font.name = "TH SarabunPSK"
                run.font.size = Pt(14)
        doc.add_paragraph()

    # สิ่งที่ต้องดำเนินการ (table)
    action_lines = sections.get("สิ่งที่ต้องทำ", [])
    parsed_actions = []
    for line in action_lines:
        if line.strip().startswith("|") and "---" not in line and "ลำดับ" not in line:
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if len(cells) >= 4:
                parsed_actions.append({"no": cells[0], "task": cells[1], "deadline": cells[2], "owner": cells[3]})
    if not parsed_actions and actions:
        for i, a in enumerate(actions):
            parsed_actions.append({"no": str(i + 1), "task": a.get("task", ""), "deadline": a.get("deadline", "TBC"), "owner": a.get("owner", "")})

    if parsed_actions:
        _add_styled_para(doc, "สิ่งที่ต้องดำเนินการต่อ", bold=True, size=16, color=HEADER_COLOR, space_after=4)

        table = doc.add_table(rows=len(parsed_actions) + 1, cols=4)
        table.style = "Table Grid"
        table.autofit = True
        headers = ["ลำดับ", "รายละเอียด", "กำหนดการ", "ผู้รับผิดชอบ"]
        widths = [Cm(1.2), Cm(9), Cm(2.8), Cm(2.8)]
        for j, h in enumerate(headers):
            _style_table_cell(table.rows[0].cells[j], h, bold=True, bg=LABEL_BG)
            table.rows[0].cells[j].width = widths[j]

        for i, item in enumerate(parsed_actions):
            _style_table_cell(table.rows[i + 1].cells[0], item["no"])
            _style_table_cell(table.rows[i + 1].cells[1], item["task"])
            _style_table_cell(table.rows[i + 1].cells[2], item["deadline"] or "TBC")
            _style_table_cell(table.rows[i + 1].cells[3], item["owner"])
            for j in range(4):
                table.rows[i + 1].cells[j].width = widths[j]

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("*หมายเหตุ : ขอความกรุณาทุกท่านตรวจสอบความถูกต้องของเอกสารสรุปการประชุมฉบับนี้")
    run.font.name = "TH SarabunPSK"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=tempfile.gettempdir())
    doc.save(tmp.name)
    tmp.close()
    return tmp.name
